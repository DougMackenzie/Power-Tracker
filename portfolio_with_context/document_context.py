"""
Document Context System
=======================
Monitors site document folders, processes new/modified files,
extracts context, and infers status changes for the tracker.

Folder Structure in Google Drive:
  Site Documents/
  ├── {site_id}/
  │   ├── meeting_notes/
  │   ├── correspondence/
  │   ├── studies/
  │   ├── contracts/
  │   └── other/
"""

import json
import os
import re
import io
import tempfile
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, field, asdict
from enum import Enum

# PDF/Doc extraction
try:
    from PyPDF2 import PdfReader
    import docx
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Google APIs
try:
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload
    GOOGLE_AVAILABLE = True
except ImportError:
    GOOGLE_AVAILABLE = False


# =============================================================================
# DATA STRUCTURES
# =============================================================================

class DocumentType(str, Enum):
    MEETING_NOTES = "meeting_notes"
    CORRESPONDENCE = "correspondence"
    STUDY = "studies"
    CONTRACT = "contracts"
    OTHER = "other"


@dataclass
class ProcessedDocument:
    """Represents a processed document with extracted content."""
    file_id: str
    file_name: str
    site_id: str
    doc_type: str
    file_path: str
    modified_time: str
    processed_time: str
    content_text: str = ""
    summary: str = ""
    key_events: List[Dict] = field(default_factory=list)
    extracted_dates: List[str] = field(default_factory=list)
    extracted_entities: List[str] = field(default_factory=list)
    status_signals: List[Dict] = field(default_factory=list)
    
    def to_dict(self) -> Dict:
        return asdict(self)
    
    @classmethod
    def from_dict(cls, data: Dict) -> 'ProcessedDocument':
        return cls(**data)


@dataclass 
class StatusChangeProposal:
    """Proposed change to tracker status based on document analysis."""
    site_id: str
    site_name: str
    field: str
    current_value: Any
    proposed_value: Any
    confidence: float  # 0-1
    evidence: str  # Quote or summary from document
    source_document: str
    source_date: str
    proposal_id: str = ""
    status: str = "pending"  # pending, approved, rejected
    
    def __post_init__(self):
        if not self.proposal_id:
            self.proposal_id = f"{self.site_id}_{self.field}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    def to_dict(self) -> Dict:
        return asdict(self)


# =============================================================================
# DOCUMENT EXTRACTION
# =============================================================================

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF bytes."""
    if not PDF_AVAILABLE:
        return "[PDF extraction not available]"
    
    try:
        reader = PdfReader(io.BytesIO(file_content))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        return f"[Error extracting PDF: {str(e)}]"


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX bytes."""
    if not PDF_AVAILABLE:
        return "[DOCX extraction not available]"
    
    try:
        doc = docx.Document(io.BytesIO(file_content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        return f"[Error extracting DOCX: {str(e)}]"


def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from plain text file."""
    try:
        return file_content.decode('utf-8', errors='ignore')
    except Exception as e:
        return f"[Error reading text: {str(e)}]"


def extract_document_text(file_content: bytes, file_name: str) -> str:
    """Extract text based on file type."""
    name_lower = file_name.lower()
    
    if name_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_content)
    elif name_lower.endswith('.docx'):
        return extract_text_from_docx(file_content)
    elif name_lower.endswith(('.txt', '.md', '.csv')):
        return extract_text_from_txt(file_content)
    elif name_lower.endswith(('.xlsx', '.xls')):
        # For Excel, return placeholder - would need pandas
        return "[Excel file - structured data extraction needed]"
    else:
        return "[Unsupported file type]"


# =============================================================================
# GOOGLE DRIVE INTEGRATION
# =============================================================================

class SiteDocumentManager:
    """Manages site document folders in Google Drive."""
    
    def __init__(self, credentials_json: str, documents_folder_id: str):
        """
        Initialize with Google credentials and root folder ID.
        
        Args:
            credentials_json: Service account credentials JSON string
            documents_folder_id: ID of "Site Documents" folder in Drive
        """
        self.documents_folder_id = documents_folder_id
        
        # Initialize Google Drive API
        creds_dict = json.loads(credentials_json)
        credentials = service_account.Credentials.from_service_account_info(
            creds_dict,
            scopes=['https://www.googleapis.com/auth/drive']
        )
        self.drive_service = build('drive', 'v3', credentials=credentials)
        
        # Cache for folder IDs
        self._folder_cache = {}
        
        # Document index (in production, store in Google Sheet or database)
        self.document_index: Dict[str, ProcessedDocument] = {}
        self.last_sync_time: Optional[datetime] = None
    
    def get_or_create_site_folder(self, site_id: str) -> str:
        """Get or create folder for a site."""
        if site_id in self._folder_cache:
            return self._folder_cache[site_id]
        
        # Check if folder exists
        query = f"name='{site_id}' and '{self.documents_folder_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
        results = self.drive_service.files().list(q=query, fields="files(id, name)").execute()
        files = results.get('files', [])
        
        if files:
            folder_id = files[0]['id']
        else:
            # Create folder
            folder_metadata = {
                'name': site_id,
                'mimeType': 'application/vnd.google-apps.folder',
                'parents': [self.documents_folder_id]
            }
            folder = self.drive_service.files().create(body=folder_metadata, fields='id').execute()
            folder_id = folder['id']
            
            # Create subfolders
            for subfolder in ['meeting_notes', 'correspondence', 'studies', 'contracts', 'other']:
                sub_metadata = {
                    'name': subfolder,
                    'mimeType': 'application/vnd.google-apps.folder',
                    'parents': [folder_id]
                }
                self.drive_service.files().create(body=sub_metadata, fields='id').execute()
        
        self._folder_cache[site_id] = folder_id
        return folder_id
    
    def list_site_folders(self) -> List[Dict]:
        """List all site folders."""
        query = f"'{self.documents_folder_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
        results = self.drive_service.files().list(
            q=query,
            fields="files(id, name, modifiedTime)"
        ).execute()
        return results.get('files', [])
    
    def list_site_documents(self, site_id: str, since: datetime = None) -> List[Dict]:
        """
        List all documents for a site, optionally filtered by modification time.
        
        Returns list of dicts with: id, name, mimeType, modifiedTime, parents
        """
        site_folder_id = self.get_or_create_site_folder(site_id)
        
        # Build query - search recursively in site folder
        query = f"'{site_folder_id}' in parents or "
        
        # Get subfolders
        subfolder_query = f"'{site_folder_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
        subfolders = self.drive_service.files().list(q=subfolder_query, fields="files(id)").execute()
        
        for sf in subfolders.get('files', []):
            query += f"'{sf['id']}' in parents or "
        
        query = query.rstrip(" or ")
        query += " and mimeType!='application/vnd.google-apps.folder' and trashed=false"
        
        if since:
            query += f" and modifiedTime > '{since.isoformat()}Z'"
        
        results = self.drive_service.files().list(
            q=query,
            fields="files(id, name, mimeType, modifiedTime, parents)",
            orderBy="modifiedTime desc"
        ).execute()
        
        return results.get('files', [])
    
    def get_document_content(self, file_id: str) -> bytes:
        """Download file content."""
        request = self.drive_service.files().get_media(fileId=file_id)
        buffer = io.BytesIO()
        downloader = MediaIoBaseDownload(buffer, request)
        
        done = False
        while not done:
            status, done = downloader.next_chunk()
        
        buffer.seek(0)
        return buffer.read()
    
    def get_document_type(self, file_path: str) -> str:
        """Determine document type from folder path."""
        path_lower = file_path.lower()
        if 'meeting' in path_lower or 'notes' in path_lower:
            return DocumentType.MEETING_NOTES
        elif 'correspond' in path_lower or 'email' in path_lower:
            return DocumentType.CORRESPONDENCE
        elif 'study' in path_lower or 'studies' in path_lower or 'sis' in path_lower or 'report' in path_lower:
            return DocumentType.STUDY
        elif 'contract' in path_lower or 'agreement' in path_lower:
            return DocumentType.CONTRACT
        else:
            return DocumentType.OTHER
    
    def get_folder_path(self, file_id: str, parents: List[str]) -> str:
        """Get folder path for a file."""
        if not parents:
            return ""
        
        try:
            parent = self.drive_service.files().get(
                fileId=parents[0],
                fields="name, parents"
            ).execute()
            return parent.get('name', '')
        except:
            return ""
    
    def scan_for_changes(self, since_hours: int = 24) -> Dict[str, List[Dict]]:
        """
        Scan all site folders for new/modified documents.
        
        Returns: {site_id: [list of changed files]}
        """
        since = datetime.utcnow() - timedelta(hours=since_hours)
        changes = {}
        
        site_folders = self.list_site_folders()
        
        for folder in site_folders:
            site_id = folder['name']
            docs = self.list_site_documents(site_id, since=since)
            
            if docs:
                changes[site_id] = docs
        
        return changes
    
    def process_document(self, file_info: Dict, site_id: str) -> ProcessedDocument:
        """
        Download and process a single document.
        
        Returns ProcessedDocument with extracted text.
        """
        file_id = file_info['id']
        file_name = file_info['name']
        
        # Get folder path for doc type detection
        parents = file_info.get('parents', [])
        folder_path = self.get_folder_path(file_id, parents)
        doc_type = self.get_document_type(folder_path)
        
        # Download and extract text
        content = self.get_document_content(file_id)
        text = extract_document_text(content, file_name)
        
        # Create processed document
        doc = ProcessedDocument(
            file_id=file_id,
            file_name=file_name,
            site_id=site_id,
            doc_type=doc_type,
            file_path=f"{site_id}/{folder_path}/{file_name}",
            modified_time=file_info.get('modifiedTime', ''),
            processed_time=datetime.utcnow().isoformat(),
            content_text=text[:50000],  # Limit to ~50K chars
        )
        
        # Store in index
        self.document_index[file_id] = doc
        
        return doc


# =============================================================================
# STATUS INFERENCE ENGINE
# =============================================================================

# Mapping of tracker fields to detection patterns and stage values
STATUS_DETECTION_CONFIG = {
    'power_stage': {
        'field_name': 'Power Confirmation',
        'stages': {
            1: {'name': 'Not Started', 'patterns': []},
            2: {'name': 'Preliminary Study', 'patterns': [
                r'preliminary\s+study',
                r'scoping\s+study',
                r'initial\s+assessment',
                r'pre-application',
            ]},
            3: {'name': 'Contract Study', 'patterns': [
                r'system\s+impact\s+study',
                r'SIS\s+(complete|submitted|results)',
                r'facilities\s+study',
                r'interconnection\s+study',
                r'contract\s+study',
            ]},
            4: {'name': 'Interconnect Agreement', 'patterns': [
                r'interconnection\s+agreement',
                r'IA\s+(signed|executed|complete)',
                r'facilities\s+agreement',
                r'FA\s+(signed|executed)',
            ]},
        }
    },
    'site_control_stage': {
        'field_name': 'Site Control',
        'stages': {
            1: {'name': 'Not Started', 'patterns': []},
            2: {'name': 'Identified', 'patterns': [
                r'site\s+identified',
                r'property\s+identified',
                r'target\s+parcel',
            ]},
            3: {'name': 'LOI', 'patterns': [
                r'letter\s+of\s+intent',
                r'LOI\s+(signed|executed|submitted)',
                r'non-binding',
            ]},
            4: {'name': 'PSA/Contract', 'patterns': [
                r'purchase\s+(agreement|contract)',
                r'PSA\s+(signed|executed)',
                r'option\s+agreement',
                r'land\s+under\s+contract',
                r'site\s+control\s+(secured|complete)',
            ]},
        }
    },
    'buyer_stage': {
        'field_name': 'Buyer Progress',
        'stages': {
            1: {'name': 'Not Started', 'patterns': []},
            2: {'name': 'Preliminary Discussion', 'patterns': [
                r'initial\s+(meeting|discussion|call)',
                r'introduced\s+to',
                r'preliminary\s+interest',
                r'exploring\s+interest',
            ]},
            3: {'name': 'LOI', 'patterns': [
                r'buyer\s+LOI',
                r'term\s+sheet',
                r'indicative\s+offer',
                r'letter\s+of\s+intent.*buyer',
            ]},
            4: {'name': 'PSA/Contract', 'patterns': [
                r'purchase\s+agreement.*buyer',
                r'buyer\s+(signed|committed|executed)',
                r'transaction\s+closed',
                r'deal\s+closed',
            ]},
        }
    },
    'zoning_stage': {
        'field_name': 'Zoning',
        'stages': {
            1: {'name': 'Not Started', 'patterns': []},
            2: {'name': 'Comp Plan/In Progress', 'patterns': [
                r'zoning\s+application',
                r'comp(rehensive)?\s+plan',
                r'planning\s+commission',
                r'zoning\s+(review|hearing|meeting)',
                r'rezoning\s+(filed|submitted|application)',
            ]},
            3: {'name': 'Zoning Approved', 'patterns': [
                r'zoning\s+(approved|complete|granted)',
                r'rezoning\s+approved',
                r'entitled',
                r'use\s+permit\s+approved',
            ]},
        }
    },
    'contract_status': {
        'field_name': 'Contract Status',
        'stages': {
            'No': {'patterns': []},
            'Verbal': {'patterns': [
                r'verbal\s+(agreement|commitment)',
                r'handshake',
                r'agreed\s+in\s+principle',
            ]},
            'MOU': {'patterns': [
                r'MOU\s+(signed|executed)',
                r'memorandum\s+of\s+understanding',
                r'letter\s+agreement',
            ]},
            'Definitive': {'patterns': [
                r'definitive\s+agreement',
                r'(signed|executed)\s+contract',
                r'binding\s+agreement',
                r'deal\s+closed',
            ]},
        }
    },
    'incentives_stage': {
        'field_name': 'Incentives',
        'stages': {
            1: {'name': 'Not Started', 'patterns': []},
            2: {'name': 'Application Filed', 'patterns': [
                r'incentive\s+application',
                r'abatement\s+application',
                r'(applied|applying)\s+for\s+incentives',
                r'TIF\s+application',
            ]},
            3: {'name': 'Preliminary Response', 'patterns': [
                r'incentive.*(preliminary|initial)\s+approval',
                r'abatement.*approved\s+in\s+principle',
            ]},
            4: {'name': 'Final Award', 'patterns': [
                r'incentive.*final\s+approval',
                r'abatement\s+(granted|approved)',
                r'incentive\s+package\s+finalized',
            ]},
        }
    },
}


def detect_status_signals(text: str, current_status: Dict) -> List[Dict]:
    """
    Analyze document text for signals that indicate status changes.
    
    Returns list of detected signals with field, proposed value, evidence, confidence.
    """
    signals = []
    text_lower = text.lower()
    
    for field, config in STATUS_DETECTION_CONFIG.items():
        current_value = current_status.get(field, 1 if field != 'contract_status' else 'No')
        
        for stage_value, stage_config in config['stages'].items():
            patterns = stage_config.get('patterns', [])
            
            for pattern in patterns:
                matches = list(re.finditer(pattern, text_lower))
                
                for match in matches:
                    # Get context around match
                    start = max(0, match.start() - 100)
                    end = min(len(text), match.end() + 100)
                    context = text[start:end]
                    
                    # Only propose if it's an advancement
                    is_advancement = False
                    if field == 'contract_status':
                        order = ['No', 'Verbal', 'MOU', 'Definitive']
                        if order.index(str(stage_value)) > order.index(str(current_value)):
                            is_advancement = True
                    else:
                        if stage_value > current_value:
                            is_advancement = True
                    
                    if is_advancement:
                        signals.append({
                            'field': field,
                            'field_name': config['field_name'],
                            'current_value': current_value,
                            'proposed_value': stage_value,
                            'stage_name': stage_config.get('name', str(stage_value)),
                            'pattern_matched': pattern,
                            'evidence': context.strip(),
                            'confidence': 0.6,  # Base confidence, will be refined by LLM
                        })
    
    # Deduplicate - keep highest stage per field
    by_field = {}
    for sig in signals:
        field = sig['field']
        if field not in by_field:
            by_field[field] = sig
        else:
            if sig['proposed_value'] > by_field[field]['proposed_value']:
                by_field[field] = sig
    
    return list(by_field.values())


# =============================================================================
# LLM-BASED STATUS INFERENCE
# =============================================================================

STATUS_INFERENCE_PROMPT = """You are analyzing a document for a data center development site to identify any status changes that should be reflected in our project tracker.

SITE INFORMATION:
- Site ID: {site_id}
- Site Name: {site_name}
- Current Status:
{current_status}

DOCUMENT INFORMATION:
- File: {file_name}
- Type: {doc_type}
- Date: {doc_date}

DOCUMENT CONTENT:
{document_text}

TRACKER FIELDS AND STAGES:

1. **power_stage** (Power Confirmation)
   - 1: Not Started
   - 2: Preliminary Study
   - 3: Contract Study (SIS, Facilities Study)
   - 4: Interconnect Agreement (IA/FA signed)

2. **site_control_stage** (Site Control)
   - 1: Not Started
   - 2: Identified
   - 3: LOI
   - 4: PSA/Contract

3. **buyer_stage** (Buyer Progress)
   - 1: Not Started
   - 2: Preliminary Discussion
   - 3: LOI
   - 4: PSA/Contract

4. **zoning_stage** (Zoning) - only 3 stages
   - 1: Not Started
   - 2: Comp Plan/In Progress
   - 3: Zoning Approved

5. **incentives_stage** (Incentives)
   - 1: Not Started
   - 2: Application Filed
   - 3: Preliminary Response
   - 4: Final Award

6. **contract_status** (Development Contract)
   - No
   - Verbal
   - MOU
   - Definitive

7. **marketing_stage** (Marketing - informational)
   - 1: Not Started
   - 2: Flyer
   - 3: VDR
   - 4: Full Package

8. **water_stage** (Water - informational)
   - 1: Not Started
   - 2: Preliminary Capacities
   - 3: Will-Serve Letter
   - 4: Final Capacities

ANALYSIS INSTRUCTIONS:
1. Read the document carefully
2. Identify any information that indicates a status change from current values
3. Only propose ADVANCEMENTS (higher stage numbers)
4. Provide specific evidence (quotes) for each proposed change
5. Assign confidence (0.0-1.0) based on how explicit the evidence is

Respond in JSON format:
{{
    "summary": "Brief summary of document content",
    "key_dates": ["list of any dates mentioned with context"],
    "key_entities": ["people, companies, agencies mentioned"],
    "proposed_changes": [
        {{
            "field": "field_name",
            "current_value": current,
            "proposed_value": new_value,
            "confidence": 0.0-1.0,
            "evidence": "Direct quote or paraphrase from document",
            "reasoning": "Why this indicates the status change"
        }}
    ],
    "action_items": ["Any follow-up actions mentioned in the document"],
    "risks_mentioned": ["Any risks or concerns noted"],
    "next_steps": ["Any next steps or upcoming milestones mentioned"]
}}

If no status changes are warranted, return an empty proposed_changes array but still provide the summary and other fields."""


def build_status_inference_prompt(
    doc: ProcessedDocument,
    site_data: Dict,
) -> str:
    """Build prompt for LLM status inference."""
    
    # Format current status
    status_lines = []
    status_fields = [
        ('power_stage', 'Power'),
        ('site_control_stage', 'Site Control'),
        ('buyer_stage', 'Buyer'),
        ('zoning_stage', 'Zoning'),
        ('incentives_stage', 'Incentives'),
        ('contract_status', 'Contract'),
        ('marketing_stage', 'Marketing'),
        ('water_stage', 'Water'),
    ]
    for field, label in status_fields:
        value = site_data.get(field, 'N/A')
        status_lines.append(f"  - {label}: {value}")
    
    return STATUS_INFERENCE_PROMPT.format(
        site_id=doc.site_id,
        site_name=site_data.get('name', doc.site_id),
        current_status="\n".join(status_lines),
        file_name=doc.file_name,
        doc_type=doc.doc_type,
        doc_date=doc.modified_time[:10] if doc.modified_time else 'Unknown',
        document_text=doc.content_text[:30000],  # Limit for context window
    )


async def infer_status_with_llm(
    doc: ProcessedDocument,
    site_data: Dict,
    llm_client: Any,
    use_claude: bool = True,
) -> Dict:
    """
    Use LLM to analyze document and infer status changes.
    
    Args:
        doc: Processed document
        site_data: Current site data from tracker
        llm_client: LLM client (Claude or Gemini)
        use_claude: If True, use Claude Opus; else use Gemini
    
    Returns:
        Dict with summary, proposed_changes, action_items, etc.
    """
    prompt = build_status_inference_prompt(doc, site_data)
    
    try:
        if use_claude:
            # Claude API
            response = llm_client.messages.create(
                model="claude-opus-4-20250514",
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}]
            )
            response_text = response.content[0].text
        else:
            # Gemini API
            response = llm_client.generate_content(prompt)
            response_text = response.text
        
        # Parse JSON response
        # Extract JSON from response (handle markdown code blocks)
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            result = json.loads(json_match.group())
            return result
        else:
            return {
                "summary": "Failed to parse response",
                "proposed_changes": [],
                "error": "No JSON found in response"
            }
    
    except Exception as e:
        return {
            "summary": f"Error during inference: {str(e)}",
            "proposed_changes": [],
            "error": str(e)
        }


def infer_status_with_llm_sync(
    doc: ProcessedDocument,
    site_data: Dict,
    llm_client: Any,
    use_claude: bool = True,
) -> Dict:
    """Synchronous version of status inference."""
    prompt = build_status_inference_prompt(doc, site_data)
    
    try:
        if use_claude:
            response = llm_client.messages.create(
                model="claude-opus-4-20250514",
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}]
            )
            response_text = response.content[0].text
        else:
            response = llm_client.generate_content(prompt)
            response_text = response.text
        
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            return json.loads(json_match.group())
        else:
            return {"summary": "Parse error", "proposed_changes": [], "error": "No JSON"}
    
    except Exception as e:
        return {"summary": f"Error: {str(e)}", "proposed_changes": [], "error": str(e)}


# =============================================================================
# CHANGE PROPOSAL MANAGEMENT
# =============================================================================

class ChangeProposalManager:
    """Manages status change proposals with approval workflow."""
    
    def __init__(self):
        self.pending_proposals: List[StatusChangeProposal] = []
        self.approved_proposals: List[StatusChangeProposal] = []
        self.rejected_proposals: List[StatusChangeProposal] = []
    
    def add_proposals_from_inference(
        self,
        inference_result: Dict,
        doc: ProcessedDocument,
        site_data: Dict,
    ) -> List[StatusChangeProposal]:
        """Create proposals from LLM inference result."""
        proposals = []
        
        for change in inference_result.get('proposed_changes', []):
            proposal = StatusChangeProposal(
                site_id=doc.site_id,
                site_name=site_data.get('name', doc.site_id),
                field=change['field'],
                current_value=change['current_value'],
                proposed_value=change['proposed_value'],
                confidence=change.get('confidence', 0.5),
                evidence=change.get('evidence', ''),
                source_document=doc.file_name,
                source_date=doc.modified_time[:10] if doc.modified_time else '',
            )
            proposals.append(proposal)
            self.pending_proposals.append(proposal)
        
        return proposals
    
    def approve_proposal(self, proposal_id: str) -> Optional[StatusChangeProposal]:
        """Mark a proposal as approved."""
        for i, p in enumerate(self.pending_proposals):
            if p.proposal_id == proposal_id:
                p.status = 'approved'
                self.approved_proposals.append(p)
                self.pending_proposals.pop(i)
                return p
        return None
    
    def reject_proposal(self, proposal_id: str) -> Optional[StatusChangeProposal]:
        """Mark a proposal as rejected."""
        for i, p in enumerate(self.pending_proposals):
            if p.proposal_id == proposal_id:
                p.status = 'rejected'
                self.rejected_proposals.append(p)
                self.pending_proposals.pop(i)
                return p
        return None
    
    def get_pending_by_site(self, site_id: str) -> List[StatusChangeProposal]:
        """Get pending proposals for a specific site."""
        return [p for p in self.pending_proposals if p.site_id == site_id]
    
    def get_all_pending(self) -> List[StatusChangeProposal]:
        """Get all pending proposals."""
        return self.pending_proposals.copy()
    
    def clear_pending(self):
        """Clear all pending proposals."""
        self.pending_proposals = []


# =============================================================================
# DOCUMENT INDEX (Simple in-memory, would use vector DB in production)
# =============================================================================

class DocumentIndex:
    """Simple document index for search and retrieval."""
    
    def __init__(self):
        self.documents: Dict[str, ProcessedDocument] = {}
        self.by_site: Dict[str, List[str]] = {}  # site_id -> [file_ids]
    
    def add_document(self, doc: ProcessedDocument):
        """Add document to index."""
        self.documents[doc.file_id] = doc
        
        if doc.site_id not in self.by_site:
            self.by_site[doc.site_id] = []
        if doc.file_id not in self.by_site[doc.site_id]:
            self.by_site[doc.site_id].append(doc.file_id)
    
    def get_site_documents(self, site_id: str) -> List[ProcessedDocument]:
        """Get all documents for a site."""
        file_ids = self.by_site.get(site_id, [])
        return [self.documents[fid] for fid in file_ids if fid in self.documents]
    
    def search_text(self, query: str, site_id: str = None) -> List[ProcessedDocument]:
        """Simple text search across documents."""
        query_lower = query.lower()
        results = []
        
        docs = self.documents.values()
        if site_id:
            docs = self.get_site_documents(site_id)
        
        for doc in docs:
            if query_lower in doc.content_text.lower():
                results.append(doc)
        
        return results
    
    def get_recent_documents(self, hours: int = 24, site_id: str = None) -> List[ProcessedDocument]:
        """Get recently modified documents."""
        cutoff = datetime.utcnow() - timedelta(hours=hours)
        cutoff_str = cutoff.isoformat()
        
        docs = self.documents.values()
        if site_id:
            docs = self.get_site_documents(site_id)
        
        return [d for d in docs if d.modified_time >= cutoff_str]


# =============================================================================
# MAIN CONTEXT MANAGER
# =============================================================================

class SiteContextManager:
    """
    Main manager for document context system.
    
    Coordinates document monitoring, processing, inference, and proposals.
    """
    
    def __init__(
        self,
        credentials_json: str,
        documents_folder_id: str,
        llm_client: Any = None,
        use_claude: bool = True,
    ):
        self.doc_manager = SiteDocumentManager(credentials_json, documents_folder_id)
        self.doc_index = DocumentIndex()
        self.proposal_manager = ChangeProposalManager()
        self.llm_client = llm_client
        self.use_claude = use_claude
        
        # Processing log
        self.last_scan_time: Optional[datetime] = None
        self.processing_log: List[Dict] = []
    
    def scan_and_process(
        self,
        sites_data: Dict[str, Dict],
        since_hours: int = 24,
    ) -> Dict:
        """
        Scan for new/modified documents, process them, and generate proposals.
        
        Args:
            sites_data: Dict of site_id -> site data (for current status)
            since_hours: Look back period for changes
        
        Returns:
            Summary of processing results
        """
        results = {
            'scan_time': datetime.utcnow().isoformat(),
            'sites_scanned': 0,
            'documents_found': 0,
            'documents_processed': 0,
            'proposals_generated': 0,
            'errors': [],
            'details': [],
        }
        
        try:
            # Scan for changes
            changes = self.doc_manager.scan_for_changes(since_hours=since_hours)
            results['sites_scanned'] = len(changes)
            
            for site_id, docs in changes.items():
                results['documents_found'] += len(docs)
                
                site_data = sites_data.get(site_id, {'name': site_id})
                
                for doc_info in docs:
                    try:
                        # Process document
                        processed = self.doc_manager.process_document(doc_info, site_id)
                        self.doc_index.add_document(processed)
                        results['documents_processed'] += 1
                        
                        # Run inference if LLM available
                        if self.llm_client and processed.content_text:
                            inference = infer_status_with_llm_sync(
                                processed,
                                site_data,
                                self.llm_client,
                                self.use_claude
                            )
                            
                            # Generate proposals
                            proposals = self.proposal_manager.add_proposals_from_inference(
                                inference,
                                processed,
                                site_data
                            )
                            results['proposals_generated'] += len(proposals)
                            
                            results['details'].append({
                                'site_id': site_id,
                                'file': doc_info['name'],
                                'summary': inference.get('summary', ''),
                                'proposals': len(proposals),
                                'action_items': inference.get('action_items', []),
                            })
                    
                    except Exception as e:
                        results['errors'].append({
                            'site_id': site_id,
                            'file': doc_info.get('name', 'unknown'),
                            'error': str(e)
                        })
            
            self.last_scan_time = datetime.utcnow()
            self.processing_log.append(results)
            
        except Exception as e:
            results['errors'].append({'error': f"Scan failed: {str(e)}"})
        
        return results
    
    def get_site_context(self, site_id: str, max_docs: int = 10) -> str:
        """
        Get combined context from all documents for a site.
        
        Useful for injecting into chat prompts.
        """
        docs = self.doc_index.get_site_documents(site_id)
        
        # Sort by modified time, most recent first
        docs.sort(key=lambda d: d.modified_time, reverse=True)
        docs = docs[:max_docs]
        
        context_parts = []
        for doc in docs:
            context_parts.append(f"""
--- Document: {doc.file_name} ({doc.doc_type}) ---
Date: {doc.modified_time[:10] if doc.modified_time else 'Unknown'}
Summary: {doc.summary or 'No summary'}
Content:
{doc.content_text[:5000]}
""")
        
        return "\n".join(context_parts)
    
    def get_pending_proposals(self, site_id: str = None) -> List[StatusChangeProposal]:
        """Get pending proposals, optionally filtered by site."""
        if site_id:
            return self.proposal_manager.get_pending_by_site(site_id)
        return self.proposal_manager.get_all_pending()
    
    def approve_proposal(self, proposal_id: str) -> Optional[Dict]:
        """Approve a proposal and return the change to apply."""
        proposal = self.proposal_manager.approve_proposal(proposal_id)
        if proposal:
            return {
                'site_id': proposal.site_id,
                'field': proposal.field,
                'value': proposal.proposed_value,
            }
        return None
    
    def reject_proposal(self, proposal_id: str) -> bool:
        """Reject a proposal."""
        return self.proposal_manager.reject_proposal(proposal_id) is not None


# =============================================================================
# EXPORTS
# =============================================================================

__all__ = [
    'DocumentType',
    'ProcessedDocument',
    'StatusChangeProposal',
    'SiteDocumentManager',
    'ChangeProposalManager',
    'DocumentIndex',
    'SiteContextManager',
    'detect_status_signals',
    'build_status_inference_prompt',
    'infer_status_with_llm_sync',
    'STATUS_DETECTION_CONFIG',
]
