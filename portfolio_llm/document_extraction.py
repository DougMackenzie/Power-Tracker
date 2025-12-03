"""
Document Extraction Module for VDR Processing
==============================================
Extracts structured site data from unstructured documents.

Supported file types:
- PDF (interconnection studies, agreements, reports)
- Word (.docx) (memos, summaries, proposals)
- Excel (.xlsx) (capacity tables, timelines, financials)
- Text/CSV (data exports, logs)

Uses pattern matching and keyword extraction to identify:
- Power capacity figures (MW)
- Study statuses (SIS, FS, FA, IA)
- Dates and timelines
- Voltage levels
- Equipment specifications
- Financial figures
"""

import re
import os
import json
import zipfile
import tempfile
from datetime import datetime, date
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, field, asdict
from pathlib import Path

# =============================================================================
# DATA STRUCTURES FOR EXTRACTED INFO
# =============================================================================

@dataclass
class ExtractedDocument:
    """Represents a processed document with extracted data."""
    filename: str
    file_type: str
    category: str  # study, agreement, financial, map, correspondence, other
    extracted_text: str
    extracted_data: Dict[str, Any] = field(default_factory=dict)
    confidence: float = 0.0
    keywords_found: List[str] = field(default_factory=list)
    processing_notes: List[str] = field(default_factory=list)


@dataclass
class VDRExtractionResult:
    """Complete result of VDR processing."""
    total_files: int = 0
    processed_files: int = 0
    failed_files: List[str] = field(default_factory=list)
    documents: List[ExtractedDocument] = field(default_factory=list)
    
    # Consolidated extracted site data
    site_name: str = ""
    utility: str = ""
    state: str = ""
    target_mw: int = 0
    
    # Power pathway
    phases: List[Dict] = field(default_factory=list)
    study_statuses: Dict[str, str] = field(default_factory=dict)
    
    # Timeline
    earliest_cod: str = ""
    phase_dates: Dict[str, str] = field(default_factory=dict)
    
    # Infrastructure
    voltages_found: List[int] = field(default_factory=list)
    transmission_miles: float = 0
    
    # Generation
    onsite_generation: List[Dict] = field(default_factory=list)
    
    # Financials
    interconnection_costs: Dict[str, float] = field(default_factory=dict)
    
    # Raw extractions for validation
    all_mw_figures: List[Tuple[str, int, str]] = field(default_factory=list)  # (context, value, source)
    all_dates: List[Tuple[str, str, str]] = field(default_factory=list)  # (context, date, source)
    all_costs: List[Tuple[str, float, str]] = field(default_factory=list)  # (context, value, source)
    
    # Validation flags
    conflicts: List[str] = field(default_factory=list)
    missing_critical: List[str] = field(default_factory=list)


# =============================================================================
# KEYWORD PATTERNS FOR EXTRACTION
# =============================================================================

# Study status patterns
STUDY_PATTERNS = {
    'sis': [
        r'system\s+impact\s+study',
        r'SIS\s+(?:complete|in\s+progress|pending|requested)',
        r'(?:completed?|finished)\s+SIS',
    ],
    'fs': [
        r'facilit(?:y|ies)\s+study',
        r'FS\s+(?:complete|in\s+progress|pending)',
        r'(?:completed?|finished)\s+(?:the\s+)?facilit(?:y|ies)\s+study',
    ],
    'fa': [
        r'facilit(?:y|ies)\s+agreement',
        r'FA\s+(?:executed|signed|pending)',
        r'(?:executed?|signed?)\s+(?:the\s+)?facilit(?:y|ies)\s+agreement',
    ],
    'ia': [
        r'interconnection\s+agreement',
        r'IA\s+(?:executed|signed|pending)',
        r'(?:executed?|signed?)\s+(?:the\s+)?interconnection\s+agreement',
        r'LGIA',
        r'GIA',
    ]
}

# Power capacity patterns
MW_PATTERNS = [
    r'(\d{1,4})\s*(?:MW|megawatt)',
    r'capacity\s+(?:of\s+)?(\d{1,4})\s*(?:MW)?',
    r'(\d{1,4})\s*MW\s+(?:of\s+)?(?:capacity|load|generation|interconnection)',
    r'phase\s*\d?\s*[:\-]?\s*(\d{1,4})\s*MW',
    r'total\s+(?:capacity|load)\s*[:\-]?\s*(\d{1,4})',
]

# Voltage patterns
VOLTAGE_PATTERNS = [
    r'(\d{2,3})\s*kV',
    r'(\d{2,3})\s*kilovolt',
    r'(\d{2,3})\s*KV',
]

# Date patterns
DATE_PATTERNS = [
    r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
    r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4})',
    r'(Q[1-4]\s+\d{4})',
    r'(\d{4})',
    r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})',
]

# Cost patterns
COST_PATTERNS = [
    r'\$\s*([\d,]+(?:\.\d{2})?)\s*(?:million|M\b)',
    r'\$\s*([\d,]+(?:\.\d{2})?)\s*(?:thousand|K\b)',
    r'\$\s*([\d,]+(?:\.\d{2})?)',
    r'([\d,]+(?:\.\d{2})?)\s*(?:million|M)\s*(?:dollars|\$)',
    r'cost[s]?\s*(?:of\s*)?\$?\s*([\d,]+)',
]

# Utility name patterns
UTILITY_PATTERNS = [
    r'(?:served\s+by|utility[:\s]+|provider[:\s]+)([A-Z][A-Za-z\s&]+(?:Power|Electric|Energy|Utility|Co\.|Company))',
    r'(PSO|AEP|Duke|Dominion|Georgia\s+Power|Southern\s+Company|Xcel|PG&E|SCE)',
    r'([A-Z][A-Za-z]+\s+(?:Power|Electric|Energy))',
]

# State patterns
STATE_PATTERNS = [
    r'\b(Oklahoma|Texas|Georgia|Virginia|Ohio|Indiana|Pennsylvania|Nevada|California|Wyoming)\b',
    r'\b(OK|TX|GA|VA|OH|IN|PA|NV|CA|WY)\b',
]

# Document category keywords
CATEGORY_KEYWORDS = {
    'study': ['system impact study', 'facilities study', 'interconnection study', 'SIS', 'FS', 'engineering study'],
    'agreement': ['agreement', 'contract', 'executed', 'signed', 'LGIA', 'GIA', 'FA', 'IA', 'PPA'],
    'financial': ['cost estimate', 'budget', 'pricing', 'invoice', 'payment', 'deposit', 'financial'],
    'map': ['site plan', 'layout', 'map', 'survey', 'plat', 'aerial'],
    'correspondence': ['email', 're:', 'fwd:', 'dear', 'regards', 'sincerely'],
    'permit': ['permit', 'zoning', 'approval', 'environmental', 'NEPA', 'Phase I', 'Phase II'],
    'technical': ['specification', 'design', 'engineering', 'voltage', 'transformer', 'substation'],
}


# =============================================================================
# FILE PARSING FUNCTIONS
# =============================================================================

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    try:
        import PyPDF2
        text = ""
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
        return text
    except ImportError:
        # Fallback: try pdfplumber
        try:
            import pdfplumber
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
            return text
        except ImportError:
            return f"[PDF extraction requires PyPDF2 or pdfplumber: {file_path}]"
    except Exception as e:
        return f"[Error extracting PDF: {str(e)}]"


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from Word document."""
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += "\n" + row_text
        
        return text
    except ImportError:
        return f"[DOCX extraction requires python-docx: {file_path}]"
    except Exception as e:
        return f"[Error extracting DOCX: {str(e)}]"


def extract_text_from_xlsx(file_path: str) -> str:
    """Extract text from Excel file."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        text = ""
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text += f"\n=== Sheet: {sheet_name} ===\n"
            
            for row in sheet.iter_rows():
                row_values = [str(cell.value) if cell.value else "" for cell in row]
                if any(row_values):
                    text += " | ".join(row_values) + "\n"
        
        return text
    except ImportError:
        return f"[XLSX extraction requires openpyxl: {file_path}]"
    except Exception as e:
        return f"[Error extracting XLSX: {str(e)}]"


def extract_text_from_file(file_path: str) -> Tuple[str, str]:
    """
    Extract text from a file based on its extension.
    Returns (text, file_type).
    """
    ext = Path(file_path).suffix.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path), 'pdf'
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path), 'docx'
    elif ext in ['.xlsx', '.xls']:
        return extract_text_from_xlsx(file_path), 'xlsx'
    elif ext in ['.txt', '.csv', '.md']:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read(), 'text'
        except Exception as e:
            return f"[Error reading text file: {str(e)}]", 'text'
    else:
        return f"[Unsupported file type: {ext}]", 'unknown'


# =============================================================================
# DATA EXTRACTION FUNCTIONS
# =============================================================================

def categorize_document(text: str, filename: str) -> str:
    """Categorize document based on content and filename."""
    text_lower = text.lower()
    filename_lower = filename.lower()
    
    scores = {cat: 0 for cat in CATEGORY_KEYWORDS}
    
    for category, keywords in CATEGORY_KEYWORDS.items():
        for keyword in keywords:
            if keyword.lower() in text_lower:
                scores[category] += 1
            if keyword.lower() in filename_lower:
                scores[category] += 2  # Filename matches weighted higher
    
    if max(scores.values()) == 0:
        return 'other'
    
    return max(scores, key=scores.get)


def extract_mw_figures(text: str, source: str) -> List[Tuple[str, int, str]]:
    """Extract MW capacity figures with context."""
    results = []
    
    for pattern in MW_PATTERNS:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            try:
                value = int(match.group(1).replace(',', ''))
                # Get surrounding context (50 chars each side)
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                context = text[start:end].replace('\n', ' ').strip()
                
                if 1 <= value <= 5000:  # Reasonable MW range
                    results.append((context, value, source))
            except (ValueError, IndexError):
                continue
    
    return results


def extract_voltages(text: str) -> List[int]:
    """Extract voltage levels from text."""
    voltages = set()
    
    for pattern in VOLTAGE_PATTERNS:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            try:
                value = int(match.group(1))
                if value in [69, 115, 138, 161, 230, 345, 500, 765]:  # Common transmission voltages
                    voltages.add(value)
            except (ValueError, IndexError):
                continue
    
    return sorted(list(voltages), reverse=True)


def extract_dates(text: str, source: str) -> List[Tuple[str, str, str]]:
    """Extract dates with context."""
    results = []
    
    for pattern in DATE_PATTERNS:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            date_str = match.group(1)
            start = max(0, match.start() - 50)
            end = min(len(text), match.end() + 50)
            context = text[start:end].replace('\n', ' ').strip()
            results.append((context, date_str, source))
    
    return results


def extract_costs(text: str, source: str) -> List[Tuple[str, float, str]]:
    """Extract cost figures with context."""
    results = []
    
    for pattern in COST_PATTERNS:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            try:
                value_str = match.group(1).replace(',', '')
                value = float(value_str)
                
                # Check for million/thousand multiplier in full match
                full_match = match.group(0).lower()
                if 'million' in full_match or full_match.endswith('m'):
                    value *= 1_000_000
                elif 'thousand' in full_match or full_match.endswith('k'):
                    value *= 1_000
                
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                context = text[start:end].replace('\n', ' ').strip()
                
                if value >= 1000:  # Filter out small numbers
                    results.append((context, value, source))
            except (ValueError, IndexError):
                continue
    
    return results


def extract_study_status(text: str) -> Dict[str, str]:
    """Extract study statuses from text."""
    statuses = {}
    text_lower = text.lower()
    
    for study_type, patterns in STUDY_PATTERNS.items():
        for pattern in patterns:
            if re.search(pattern, text_lower):
                # Try to determine status
                if 'complete' in text_lower or 'finished' in text_lower:
                    statuses[study_type] = 'complete'
                elif 'in progress' in text_lower or 'ongoing' in text_lower:
                    statuses[study_type] = 'in_progress'
                elif 'executed' in text_lower or 'signed' in text_lower:
                    statuses[study_type] = 'executed'
                elif 'pending' in text_lower or 'requested' in text_lower:
                    statuses[study_type] = 'requested'
                else:
                    statuses[study_type] = 'mentioned'
                break
    
    return statuses


def extract_utility_name(text: str) -> Optional[str]:
    """Extract utility name from text."""
    for pattern in UTILITY_PATTERNS:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return None


def extract_state(text: str) -> Optional[str]:
    """Extract state from text."""
    # First try full names
    for pattern in STATE_PATTERNS:
        match = re.search(pattern, text)
        if match:
            state = match.group(1)
            # Normalize to code
            state_map = {
                'Oklahoma': 'OK', 'Texas': 'TX', 'Georgia': 'GA',
                'Virginia': 'VA', 'Ohio': 'OH', 'Indiana': 'IN',
                'Pennsylvania': 'PA', 'Nevada': 'NV', 'California': 'CA',
                'Wyoming': 'WY'
            }
            return state_map.get(state, state)
    return None


# =============================================================================
# MAIN VDR PROCESSING
# =============================================================================

def process_vdr_upload(file_path: str) -> VDRExtractionResult:
    """
    Process a VDR upload (zip file or directory).
    Returns consolidated extraction result.
    """
    result = VDRExtractionResult()
    
    # Handle zip file
    if file_path.endswith('.zip'):
        with tempfile.TemporaryDirectory() as temp_dir:
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            result = _process_directory(temp_dir, result)
    elif os.path.isdir(file_path):
        result = _process_directory(file_path, result)
    else:
        # Single file
        result = _process_single_file(file_path, result)
    
    # Consolidate and validate
    result = _consolidate_extractions(result)
    result = _validate_extractions(result)
    
    return result


def _process_directory(dir_path: str, result: VDRExtractionResult) -> VDRExtractionResult:
    """Process all files in a directory recursively."""
    
    for root, dirs, files in os.walk(dir_path):
        # Skip hidden directories
        dirs[:] = [d for d in dirs if not d.startswith('.')]
        
        for filename in files:
            if filename.startswith('.'):
                continue
            
            file_path = os.path.join(root, filename)
            result.total_files += 1
            
            try:
                result = _process_single_file(file_path, result)
                result.processed_files += 1
            except Exception as e:
                result.failed_files.append(f"{filename}: {str(e)}")
    
    return result


def _process_single_file(file_path: str, result: VDRExtractionResult) -> VDRExtractionResult:
    """Process a single file and add to result."""
    filename = os.path.basename(file_path)
    
    # Extract text
    text, file_type = extract_text_from_file(file_path)
    
    if file_type == 'unknown' or text.startswith('['):
        result.failed_files.append(filename)
        return result
    
    # Categorize document
    category = categorize_document(text, filename)
    
    # Extract data
    extracted_data = {}
    keywords_found = []
    
    # MW figures
    mw_figures = extract_mw_figures(text, filename)
    if mw_figures:
        extracted_data['mw_figures'] = mw_figures
        result.all_mw_figures.extend(mw_figures)
        keywords_found.append('MW capacity')
    
    # Voltages
    voltages = extract_voltages(text)
    if voltages:
        extracted_data['voltages'] = voltages
        result.voltages_found.extend(voltages)
        keywords_found.append('voltage')
    
    # Dates
    dates = extract_dates(text, filename)
    if dates:
        extracted_data['dates'] = dates
        result.all_dates.extend(dates)
        keywords_found.append('dates')
    
    # Costs
    costs = extract_costs(text, filename)
    if costs:
        extracted_data['costs'] = costs
        result.all_costs.extend(costs)
        keywords_found.append('costs')
    
    # Study statuses
    study_statuses = extract_study_status(text)
    if study_statuses:
        extracted_data['study_statuses'] = study_statuses
        result.study_statuses.update(study_statuses)
        keywords_found.append('study status')
    
    # Utility
    utility = extract_utility_name(text)
    if utility:
        extracted_data['utility'] = utility
        if not result.utility:
            result.utility = utility
        keywords_found.append('utility')
    
    # State
    state = extract_state(text)
    if state:
        extracted_data['state'] = state
        if not result.state:
            result.state = state
        keywords_found.append('state')
    
    # Calculate confidence
    confidence = min(len(keywords_found) / 5, 1.0)
    
    # Create document record
    doc = ExtractedDocument(
        filename=filename,
        file_type=file_type,
        category=category,
        extracted_text=text[:5000],  # Truncate for storage
        extracted_data=extracted_data,
        confidence=confidence,
        keywords_found=keywords_found
    )
    
    result.documents.append(doc)
    
    return result


def _consolidate_extractions(result: VDRExtractionResult) -> VDRExtractionResult:
    """Consolidate extracted data into site fields."""
    
    # Consolidate MW figures - take the most common large value
    if result.all_mw_figures:
        mw_values = [v for _, v, _ in result.all_mw_figures if v >= 100]
        if mw_values:
            # Use the largest value as target
            result.target_mw = max(mw_values)
    
    # Consolidate voltages
    result.voltages_found = sorted(list(set(result.voltages_found)), reverse=True)
    
    return result


def _validate_extractions(result: VDRExtractionResult) -> VDRExtractionResult:
    """Validate extractions and flag conflicts/missing data."""
    
    # Check for conflicting MW figures
    if result.all_mw_figures:
        unique_mw = set(v for _, v, _ in result.all_mw_figures if v >= 100)
        if len(unique_mw) > 3:
            result.conflicts.append(f"Multiple MW figures found: {sorted(unique_mw)}")
    
    # Check for missing critical data
    if not result.utility:
        result.missing_critical.append("Utility name not found")
    if not result.state:
        result.missing_critical.append("State not identified")
    if result.target_mw == 0:
        result.missing_critical.append("No MW capacity figures found")
    if not result.study_statuses:
        result.missing_critical.append("No study status information found")
    
    return result


def extraction_result_to_site_data(result: VDRExtractionResult) -> Dict:
    """Convert extraction result to site database format."""
    
    # Map study statuses to app format
    study_status_map = {
        'mentioned': 'sis_requested',
        'requested': 'sis_requested',
        'in_progress': 'sis_in_progress',
        'complete': 'sis_complete',
        'executed': 'fa_executed'
    }
    
    # Determine highest study status
    study_status = 'not_started'
    study_priority = ['ia', 'fa', 'fs', 'sis']
    
    for study in study_priority:
        if study in result.study_statuses:
            status = result.study_statuses[study]
            if study == 'ia' and status == 'executed':
                study_status = 'ia_executed'
                break
            elif study == 'fa' and status == 'executed':
                study_status = 'fa_executed'
                break
            elif study == 'fs' and status == 'complete':
                study_status = 'fs_complete'
                break
            elif study == 'sis':
                study_status = study_status_map.get(status, 'sis_requested')
                break
    
    return {
        'name': result.site_name or "Imported Site",
        'state': result.state or '',
        'utility': result.utility or '',
        'target_mw': result.target_mw,
        'study_status': study_status,
        'notes': f"Imported from VDR. Files processed: {result.processed_files}. Voltages: {result.voltages_found}",
        
        # Fields requiring validation
        '_validation_needed': True,
        '_mw_figures_found': [(c, v, s) for c, v, s in result.all_mw_figures],
        '_conflicts': result.conflicts,
        '_missing': result.missing_critical,
    }


# =============================================================================
# CONVERSATIONAL EXTRACTION (For Chat Interface)
# =============================================================================

def generate_extraction_prompt(user_description: str) -> str:
    """Generate a prompt for Claude to extract structured data from natural language."""
    
    return f"""Extract structured site information from this description. Return a JSON object with the following fields (use null for unknown values):

{{
    "site_name": "string or null",
    "state": "2-letter code or null",
    "utility": "utility company name or null",
    "target_mw": "integer or null",
    "acreage": "integer or null",
    "study_status": "not_started|sis_requested|sis_in_progress|sis_complete|fs_complete|fa_executed|ia_executed",
    "utility_commitment": "none|initial|engaged|verbal|committed",
    "power_timeline_months": "integer or null",
    "voltages_mentioned": ["list of kV values"],
    "transmission_miles": "float or null",
    "land_control": "none|negotiating|loi|option|owned",
    "water_status": "unknown|constrained|identified|available|secured",
    "zoning_status": "unknown|not_started|in_progress|approved",
    "onsite_generation": [
        {{"type": "gas|solar|battery|smr", "capacity_mw": int, "status": "string"}}
    ],
    "end_user_interest": "none|interest_expressed|tours_completed|nda_signed|nda_active|loi|term_sheet",
    "key_risks": ["list of identified risks"],
    "key_opportunities": ["list of opportunities"],
    "additional_notes": "any other relevant information"
}}

User description:
{user_description}

JSON response:"""


def parse_conversational_input(user_text: str) -> Dict:
    """
    Parse natural language site description into structured fields.
    This is a rule-based fallback; ideally use Claude API for better extraction.
    """
    
    result = {
        'site_name': None,
        'state': None,
        'utility': None,
        'target_mw': None,
        'study_status': 'not_started',
        'utility_commitment': 'none',
        'voltages_mentioned': [],
        'land_control': 'none',
        'water_status': 'unknown',
        'extracted_keywords': []
    }
    
    text = user_text.lower()
    
    # Extract state
    state = extract_state(user_text)
    if state:
        result['state'] = state
        result['extracted_keywords'].append(f"State: {state}")
    
    # Extract utility
    utility = extract_utility_name(user_text)
    if utility:
        result['utility'] = utility
        result['extracted_keywords'].append(f"Utility: {utility}")
    
    # Extract MW
    mw_figures = extract_mw_figures(user_text, "user_input")
    if mw_figures:
        result['target_mw'] = max(v for _, v, _ in mw_figures)
        result['extracted_keywords'].append(f"Capacity: {result['target_mw']} MW")
    
    # Extract voltages
    voltages = extract_voltages(user_text)
    if voltages:
        result['voltages_mentioned'] = voltages
        result['extracted_keywords'].append(f"Voltages: {voltages}")
    
    # Study status keywords
    if 'ia executed' in text or 'interconnection agreement signed' in text:
        result['study_status'] = 'ia_executed'
    elif 'fa executed' in text or 'facilities agreement signed' in text:
        result['study_status'] = 'fa_executed'
    elif 'facilities study complete' in text or 'fs complete' in text:
        result['study_status'] = 'fs_complete'
    elif 'sis complete' in text or 'system impact study complete' in text:
        result['study_status'] = 'sis_complete'
    elif 'sis in progress' in text or 'studying' in text:
        result['study_status'] = 'sis_in_progress'
    elif 'sis requested' in text or 'submitted' in text:
        result['study_status'] = 'sis_requested'
    
    # Utility commitment keywords
    if 'committed' in text or 'commitment letter' in text:
        result['utility_commitment'] = 'committed'
    elif 'verbal' in text or 'handshake' in text:
        result['utility_commitment'] = 'verbal'
    elif 'engaged' in text or 'discussing' in text or 'working with' in text:
        result['utility_commitment'] = 'engaged'
    elif 'initial' in text or 'early' in text:
        result['utility_commitment'] = 'initial'
    
    # Land control
    if 'own' in text or 'purchased' in text:
        result['land_control'] = 'owned'
    elif 'option' in text:
        result['land_control'] = 'option'
    elif 'loi' in text or 'letter of intent' in text:
        result['land_control'] = 'loi'
    elif 'negotiat' in text:
        result['land_control'] = 'negotiating'
    
    # Water
    if 'water secured' in text or 'water rights' in text:
        result['water_status'] = 'secured'
    elif 'water available' in text or 'municipal water' in text:
        result['water_status'] = 'available'
    elif 'water identified' in text:
        result['water_status'] = 'identified'
    elif 'water constrained' in text or 'water issues' in text:
        result['water_status'] = 'constrained'
    
    return result


if __name__ == "__main__":
    # Test conversational parsing
    test_description = """
    We're looking at a 500MW site in Oklahoma, served by PSO. 
    The SIS is complete and we're waiting on the facilities study. 
    It's adjacent to a 230kV line about 3 miles away. 
    We have the land under option and water is available from the municipal system.
    PSO has been engaged and seems supportive.
    """
    
    print("Testing conversational extraction:")
    print("-" * 50)
    result = parse_conversational_input(test_description)
    print(json.dumps(result, indent=2))
